#testdocx.py
from docx import Document #จากไลบราลี่ docx เรียกคลาส document
import wikipedia

def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    #summary สำหรับบทความที่สรุป
    data=wikipedia.summary(keyword)

    #page+content บทความทั้งหน้า
    data2=wikipedia.page(keyword)
    data2=data2.content

    doc=Document() #สร้างไฟล์ word ในไพธอน
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword+'.docx')
    print('สร้างไฟล์สำเร็จ')

try:
    Wiki('dddddประเทศไทย')
except:
    print('ไม่พบ')
    
#Wiki('United States','en')
#Wiki('ประเทศญี่ปุ่น','jp')



