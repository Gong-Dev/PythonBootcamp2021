#GuiWiki.py
import wikipedia

#python to docx
from docx import Document
def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    #summary สำหรับบทความที่สรุป
    data=wikipedia.summary(keyword)

    #page+content บทความทั้งหน้า
    data2=wikipedia.page(keyword)
    data2=data2.content

    doc=Document() #สร้างไฟล์ word ในไพธอน
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword+'.docx')
    print('สร้างไฟล์สำเร็จ')

#เปลี่ยนเป็นภาษาไทย
wikipedia.set_lang('th') 
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

#config
Font1=('Angsana New',20)


#สร้างGUI
GUI=Tk()
GUI.title('wiki by GONGGANG')
GUI.geometry('400x300')
#คำอธิบาย
L=ttk.Label(GUI,text='ค้นหาบทความ',font=Font1)
#ช่องค้นหาข้อมูล
v_search=StringVar()
E1=ttk.Entry(GUI,textvariable=v_search,font=Font1,width=40) #สร้างช่องค้นหาข้อมูล
E1.pack(pady=10)

#ปุ่มค้นหา
def Search():
    keyword=v_search.get() #.get()คือดึงข้อมูลเข้ามา
    try :
        #ลองค้นหาดูว่าได้ผลลัพธ์หรือไม่ หากได้ให้ผ่านไป
        language=v_radio.get() #th/en/zh
        Wiki(keyword)
        messagebox.showinfo('บันทึกสำเร็จ','ค้นหาข้อความสำเร็จและบันทึกแล้ว')
    except:
        #หากรันคำสั่งแล้วมีปัญหา แสดงข้อความแจ้งเตือน
        messagebox.showwarning('Keyword Error','กรุณากรอกคำค้นหาใหม่')

    #print (wikipedia.search(keyword))
    #result=wikipedia.summary(keyword)
    #print(result)



B1=ttk.Button(GUI,text='ค้นหา',command=Search)
B1.pack(ipadx=20,ipady=10) #ipadx ขยายภายในปุ่มแนวนอน

#เลือกภาษา
f1=Frame(GUI)
f1.pack(pady=20)

v_radio=StringVar() #ช่องเก็บภาษา

rb1=ttk.Radiobutton(f1,text='ไทย',variable=v_radio,value='th')
rb2=ttk.Radiobutton(f1,text='อังกฤษ',variable=v_radio,value='en')
rb3=ttk.Radiobutton(f1,text='จีน',variable=v_radio,value='zh')
rb1.invoke() #สั่งให้ค่าเริ่มเป็นภาษาไทย

rb1.grid(row=0,column=0)
rb2.grid(row=0,column=1)
rb3.grid(row=0,column=2)

GUI.mainloop()


